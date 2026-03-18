import re
from typing import Iterable, List, Optional

from docx import Document

from src.backend.models import Requirement, TestCase


REQ_PREFIXES_SYSTEM = ("BNC", "DLS", "NSE")
REQ_PREFIXES_STAKEHOLDER = ("BNCS", "DLSS", "NSES")

ID_PATTERN = re.compile(
    r"[#\[\(]?\s*(?P<id>(?:BNCS|DLSS|NSES|BNC|DLS|NSE)\s*-?\s*\d{3,})\s*[\]\)]?"
)
TS_PATTERN = re.compile(r"^\d+(?:\.\d+)*$")
TEST_CASE_PATTERN = re.compile(
    r"^\s*(?P<section>\d+(?:\.\d+)*)\s+test\s*case\s*(?P<tcnum>\d+)?\s*:?\s*(?P<title>.*)$",
    re.IGNORECASE,
)
BODY_TEST_CASE_PATTERN = re.compile(
    r"^\s*test\s*case\s*:?\s*(?P<title>.+?)\s*$",
    re.IGNORECASE,
)


def _is_section(text: str, targets: Iterable[str]) -> bool:
    normalized = " ".join((text or "").strip().lower().split())
    if normalized in targets:
        return True
    stripped = re.sub(r"^[\d\.]+\s*", "", normalized)
    return stripped in targets


def _extract_ids(text: str) -> List[str]:
    ids = []
    for m in ID_PATTERN.finditer(text or ""):
        raw = m.group("id")
        cleaned = re.sub(r"[\s-]+", "", raw or "")
        if cleaned:
            ids.append(cleaned)
    return ids


def _normalize_header(text: str) -> str:
    return re.sub(r"[^a-z0-9]", "", (text or "").lower())


def _normalize_title(text: str) -> str:
    return " ".join((text or "").strip().lower().split())


def _title_lookup_key(text: str) -> str:
    normalized = _normalize_title(text)
    return re.sub(r"\s+\d+$", "", normalized)


def _find_header_indices(header_row: List[str]) -> Optional[tuple]:
    normalized = [_normalize_header(c) for c in header_row]
    ts_idx = None
    ref_idx = None
    for i, cell in enumerate(normalized):
        if cell == "ts":
            ts_idx = i
        if cell in ("ref", "reference"):
            ref_idx = i
    if ts_idx is None or ref_idx is None:
        return None
    return ts_idx, ref_idx


def _is_acceptance_heading(text: str) -> bool:
    normalized = " ".join(text.strip().lower().split())
    return normalized in ("accept criteria", "acceptance criteria")


def parse_requirements(doc_path: str, source_label: Optional[str] = None) -> List[Requirement]:
    doc = Document(doc_path)
    requirements: List[Requirement] = []
    seen = set()
    current_stakeholder: Optional[str] = None
    source_doc = source_label or doc_path
    source_doc = source_label or doc_path

    def add_system_ids(ids: List[str], stakeholder: Optional[str]) -> None:
        system_ids = [i for i in ids if i.startswith(REQ_PREFIXES_SYSTEM) and not i.startswith(REQ_PREFIXES_STAKEHOLDER)]
        for req_id in system_ids:
            key = (req_id, stakeholder, doc_path)
            if key in seen:
                continue
            seen.add(key)
            requirements.append(
                Requirement(req_id=req_id, stakeholder_id=stakeholder, source_doc=source_doc)
            )

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue

        ids = _extract_ids(text)
        stakeholder_ids = [i for i in ids if i.startswith(REQ_PREFIXES_STAKEHOLDER)]
        if stakeholder_ids:
            current_stakeholder = stakeholder_ids[-1]
        add_system_ids(ids, current_stakeholder)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = (cell.text or "").strip()
                if not text:
                    continue
                ids = _extract_ids(text)
                stakeholder_ids = [i for i in ids if i.startswith(REQ_PREFIXES_STAKEHOLDER)]
                if stakeholder_ids:
                    current_stakeholder = stakeholder_ids[-1]
                add_system_ids(ids, current_stakeholder)

    return requirements


def parse_test_procedures(doc_path: str, source_label: Optional[str] = None) -> List[TestCase]:
    doc = Document(doc_path)
    test_cases: List[TestCase] = []
    source_doc = source_label or doc_path

    tc_requirements = {}
    tc_titles = {}
    title_to_tc = {}
    current_tc: Optional[str] = None
    in_requirements = False
    has_body_test_case_headings = any(
        BODY_TEST_CASE_PATTERN.match((p.text or "").strip()) for p in doc.paragraphs
    )
    in_body_scope = not has_body_test_case_headings

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue

        tc_match = TEST_CASE_PATTERN.match(text)
        if tc_match:
            tc_num = tc_match.group("tcnum") or tc_match.group("section")
            current_tc = tc_num
            title = (tc_match.group("title") or "").strip()
            if title:
                tc_titles[tc_num] = title
                normalized_title = _title_lookup_key(title)
                if normalized_title:
                    title_to_tc[normalized_title] = tc_num
            in_requirements = False
            if has_body_test_case_headings:
                # In documents with explicit "Test Case : <title>" body headings,
                # numbered headings are typically TOC entries and should not set scope.
                current_tc = None
            continue

        body_tc_match = BODY_TEST_CASE_PATTERN.match(text)
        if body_tc_match:
            title = (body_tc_match.group("title") or "").strip()
            mapped_tc = title_to_tc.get(_title_lookup_key(title))
            current_tc = mapped_tc
            in_requirements = False
            in_body_scope = True
            continue

        if not in_body_scope:
            continue

        if _is_section(text, {"requirements", "requirement"}):
            in_requirements = current_tc is not None
            continue

        if _is_section(
            text,
            {
                "test execution",
                "test steps",
                "test procedure",
                "test execution steps",
                "comments",
            },
        ):
            in_requirements = False
            continue

        if in_requirements and current_tc:
            ids = _extract_ids(text)
            if ids:
                tc_requirements.setdefault(current_tc, set()).update(ids)

    for tc_id, ids in tc_requirements.items():
        title = tc_titles.get(tc_id)
        for ref_id in sorted(ids):
            test_cases.append(
                TestCase(
                    ts_id="",
                    ref_id=ref_id,
                    source_doc=source_doc,
                    test_case_id=tc_id,
                    test_case_title=title,
                )
            )

    for table in doc.tables:
        header_idx = None
        header_cols = None
        for i, row in enumerate(table.rows):
            row_text = [c.text.strip() for c in row.cells]
            header_cols = _find_header_indices(row_text)
            if header_cols is not None:
                header_idx = i
                break

        if header_idx is None or header_cols is None:
            continue

        ts_col, ref_col = header_cols
        for row in table.rows[header_idx + 1 :]:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < max(ts_col, ref_col) + 1:
                continue

            ts_val = cells[ts_col]
            ref_val = cells[ref_col]

            if not ts_val or not ref_val:
                continue

            if not TS_PATTERN.match(ts_val):
                continue

            ids = _extract_ids(ref_val)
            if not ids:
                continue
            tc_id = ts_val.split(".", 1)[0]
            title = tc_titles.get(tc_id)
            for ref_id in ids:
                test_cases.append(
                    TestCase(
                        ts_id=ts_val,
                        ref_id=ref_id,
                        source_doc=source_doc,
                        test_case_id=tc_id,
                        test_case_title=title,
                    )
                )

    return test_cases
